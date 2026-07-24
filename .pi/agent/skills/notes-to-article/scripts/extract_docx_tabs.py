#!/usr/bin/env python3
"""Extract Tab 1/Tab 2-style sections from a DOCX file."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: python-docx. Install it in the active Python environment."
    ) from exc


MARKERS = {
    "tab_1": {"tab 1", "tab1", "source", "input", "notes"},
    "tab_2": {"tab 2", "tab2", "output", "article", "reference"},
}


def iter_docx_blocks(doc: Document) -> list[str]:
    blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append("\n".join(cells))
    return blocks


def extract_docx(path: Path) -> dict[str, object]:
    doc = Document(path)
    blocks = iter_docx_blocks(doc)

    sections: dict[str, list[str]] = {"tab_1": [], "tab_2": []}
    current: str | None = None
    loose: list[str] = []

    for text in blocks:
        normalized = text.lower().strip()
        if normalized in MARKERS["tab_1"]:
            current = "tab_1"
            continue
        if normalized in MARKERS["tab_2"]:
            current = "tab_2"
            continue
        if current:
            sections[current].append(text)
        else:
            loose.append(text)

    return {
        "path": str(path),
        "paragraph_count": len(blocks),
        "loose": loose,
        "tab_1": "\n".join(sections["tab_1"]),
        "tab_2": "\n".join(sections["tab_2"]),
        "has_tab_1": bool(sections["tab_1"]),
        "has_tab_2": bool(sections["tab_2"]),
    }


def selected_text(result: dict[str, object], section: str | None) -> str:
    if section:
        return str(result[section])

    parts: list[str] = []
    if result["loose"]:
        parts.append("## loose\n\n" + "\n".join(result["loose"]))  # type: ignore[arg-type]
    for key in ("tab_1", "tab_2"):
        value = str(result[key])
        if value:
            parts.append(f"## {key}\n\n{value}")
    return "\n\n".join(parts)


def write_docx(path: Path, text: str) -> None:
    doc = Document()
    for block in text.split("\n"):
        doc.add_paragraph(block)
    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--json", action="store_true", help="Print JSON instead of text.")
    parser.add_argument("--section", choices=("tab_1", "tab_2"), help="Print or write one section.")
    parser.add_argument("--output", type=Path, help="Write selected text to a UTF-8 .txt file.")
    parser.add_argument("--output-docx", type=Path, help="Write selected text to a .docx file.")
    args = parser.parse_args()

    result = extract_docx(args.docx)
    text = selected_text(result, args.section)

    if args.output:
        args.output.write_text(text, encoding="utf-8")

    if args.output_docx:
        write_docx(args.output_docx, text)

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif not args.output and not args.output_docx:
        print(text)

    return 0


if __name__ == "__main__":
    sys.exit(main())
